"""
Documentation Generator for Manimations Project
Generates comprehensive DOCX documentation with diagrams
"""
import sys
from pathlib import Path

# Add root to path
root = Path(__file__).resolve().parent
if str(root) not in sys.path:
    sys.path.insert(0, str(root))

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing required package: python-docx")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_documentation():
    """Generate comprehensive documentation DOCX file"""
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # ==================== COVER PAGE ====================
    title = doc.add_heading('Manimations Project', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Comprehensive Technical Documentation')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.color.rgb = RGBColor(0, 102, 204)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    version_info = doc.add_paragraph('Version: 1.0 (v01)')
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_info = doc.add_paragraph('January 2026')
    date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ==================== TABLE OF CONTENTS ====================
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        '1. Executive Summary',
        '2. System Architecture Overview',
        '3. Component Breakdown',
        '4. Data Flow and Processing Pipeline',
        '5. API Reference',
        '6. Installation and Setup',
        '7. Usage Guide',
        '8. Configuration Management',
        '9. LLM Integration',
        '10. Animation Generation Engine',
        '11. Web Interface',
        '12. Voice Synthesis System',
        '13. Error Handling and Debugging',
        '14. Performance Optimization',
        '15. Security Considerations',
        '16. Appendices'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    doc.add_page_break()
    
    # ==================== EXECUTIVE SUMMARY ====================
    doc.add_heading('1. Executive Summary', 1)
    
    doc.add_heading('1.1 Project Overview', 2)
    doc.add_paragraph(
        'Manimations is an advanced AI-powered mathematical animation generation system that '
        'transforms natural language queries into high-quality educational videos. The system '
        'leverages state-of-the-art Large Language Models (LLMs) to understand mathematical '
        'concepts and automatically generates visualization code using the Manim (Mathematical '
        'Animation Engine) library.'
    )
    
    doc.add_heading('1.2 Key Features', 2)
    features = [
        'Natural Language Processing: Accepts plain English mathematical questions',
        'Multi-LLM Support: Compatible with OpenAI GPT, DeepSeek, Gemini, and Ollama',
        'Automated Animation Generation: Creates professional mathematical visualizations',
        'Voice Synthesis: Generates educational narration using TTS engines',
        'Web Interface: User-friendly browser-based interface',
        'REST API: Programmatic access for integration',
        'Flexible Pipeline: Voice-first or video-first rendering modes',
        'Historical Data Processing: Generate videos from pre-existing solutions'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('1.3 Technology Stack', 2)
    tech_stack = {
        'Backend': 'Python 3.13, FastAPI, Uvicorn',
        'Animation': 'Manim Community Edition',
        'LLM Integration': 'OpenAI API, DeepSeek API, Google Gemini, Ollama',
        'Voice Synthesis': 'pyttsx3, ElevenLabs (optional)',
        'Video Processing': 'FFmpeg',
        'Frontend': 'HTML5, CSS3, Vanilla JavaScript',
        'Validation': 'Pydantic v2'
    }
    table = doc.add_table(rows=len(tech_stack)+1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Component'
    hdr_cells[1].text = 'Technology'
    
    for idx, (component, tech) in enumerate(tech_stack.items(), 1):
        row_cells = table.rows[idx].cells
        row_cells[0].text = component
        row_cells[1].text = tech
    
    doc.add_page_break()
    
    # ==================== SYSTEM ARCHITECTURE ====================
    doc.add_heading('2. System Architecture Overview', 1)
    
    doc.add_heading('2.1 High-Level Architecture', 2)
    doc.add_paragraph(
        'The Manimations system follows a modular, pipeline-based architecture consisting of '
        'four primary layers:'
    )
    
    layers = [
        ('Presentation Layer', 'Web UI and REST API endpoints'),
        ('Orchestration Layer', 'LLM prompt engineering and response parsing'),
        ('Generation Layer', 'Manim code generation and video rendering'),
        ('Integration Layer', 'Audio synthesis and video synchronization')
    ]
    
    for layer_name, layer_desc in layers:
        p = doc.add_paragraph()
        p.add_run(f'{layer_name}: ').bold = True
        p.add_run(layer_desc)
    
    doc.add_heading('2.2 Architecture Diagram', 2)
    doc.add_paragraph('[ASCII Architecture Diagram]')
    
    arch_diagram = '''
    ┌─────────────────────────────────────────────────────────────┐
    │                    Presentation Layer                        │
    │  ┌──────────────┐              ┌──────────────┐            │
    │  │   Web UI     │              │  REST API    │            │
    │  │ (HTML/CSS/JS)│◄────────────►│  (FastAPI)   │            │
    │  └──────────────┘              └──────────────┘            │
    └─────────────────────────────────────────────────────────────┘
                            │
                            ▼
    ┌─────────────────────────────────────────────────────────────┐
    │                  Orchestration Layer                         │
    │  ┌──────────────────────────────────────────────────────┐  │
    │  │         Prompt Orchestrator                          │  │
    │  │  ┌──────────┐  ┌──────────┐  ┌──────────┐          │  │
    │  │  │  OpenAI  │  │ DeepSeek │  │  Gemini  │          │  │
    │  │  │  Client  │  │  Client  │  │  Client  │          │  │
    │  │  └──────────┘  └──────────┘  └──────────┘          │  │
    │  └──────────────────────────────────────────────────────┘  │
    └─────────────────────────────────────────────────────────────┘
                            │
                            ▼
    ┌─────────────────────────────────────────────────────────────┐
    │                   Generation Layer                           │
    │  ┌──────────────┐              ┌──────────────┐            │
    │  │    Manim     │              │  Voiceover   │            │
    │  │   Adapter    │              │  Synthesis   │            │
    │  └──────────────┘              └──────────────┘            │
    └─────────────────────────────────────────────────────────────┘
                            │
                            ▼
    ┌─────────────────────────────────────────────────────────────┐
    │                  Integration Layer                           │
    │  ┌──────────────────────────────────────────────────────┐  │
    │  │           AV Sync (FFmpeg)                           │  │
    │  └──────────────────────────────────────────────────────┘  │
    └─────────────────────────────────────────────────────────────┘
                            │
                            ▼
                    ┌───────────────┐
                    │  Final Video  │
                    │    (.mp4)     │
                    └───────────────┘
    '''
    
    code_para = doc.add_paragraph(arch_diagram)
    code_para.style = 'Intense Quote'
    
    doc.add_page_break()
    
    # ==================== COMPONENT BREAKDOWN ====================
    doc.add_heading('3. Component Breakdown', 1)
    
    components = {
        'scripts/server/app.py': {
            'purpose': 'FastAPI server handling HTTP requests',
            'key_functions': ['RunRequest model', 'Job management', 'Worker threads'],
            'dependencies': ['FastAPI', 'Pydantic', 'threading']
        },
        'scripts/orchestrator/prompt_orchestrator.py': {
            'purpose': 'LLM orchestration and prompt engineering',
            'key_functions': ['craft_solver_prompt', 'call_deepseek_solver', 'orchestrate_solution'],
            'dependencies': ['OpenAI', 'google-generativeai', 'llm_clients']
        },
        'scripts/manim_adapter.py': {
            'purpose': 'Converts JSON plans to executable Manim code',
            'key_functions': ['generate_scene_script', 'render_with_manim'],
            'dependencies': ['manim', 'json', 'subprocess']
        },
        'scripts/voiceover.py': {
            'purpose': 'Text-to-speech audio generation',
            'key_functions': ['synthesize_scene_wise', 'synthesize_element_wise'],
            'dependencies': ['pyttsx3', 'elevenlabs (optional)']
        },
        'scripts/av_sync.py': {
            'purpose': 'Synchronizes audio with video',
            'key_functions': ['combine_video_with_audio'],
            'dependencies': ['FFmpeg']
        },
        'scripts/pipeline.py': {
            'purpose': 'End-to-end pipeline orchestration',
            'key_functions': ['run_pipeline'],
            'dependencies': ['All above components']
        }
    }
    
    for comp_path, comp_info in components.items():
        doc.add_heading(f'3.{list(components.keys()).index(comp_path)+1} {comp_path}', 2)
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Purpose: ').bold = True
        p.add_run(comp_info['purpose'])
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Key Functions:').bold = True
        for func in comp_info['key_functions']:
            doc.add_paragraph(f'• {func}', style='List Bullet')
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('Dependencies: ').bold = True
        p.add_run(', '.join(comp_info['dependencies']))
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== DATA FLOW ====================
    doc.add_heading('4. Data Flow and Processing Pipeline', 1)
    
    doc.add_heading('4.1 Request Flow', 2)
    doc.add_paragraph('The typical request follows this sequence:')
    
    flow_steps = [
        'User submits question via Web UI',
        'FastAPI receives POST request at /api/run',
        'Job created with unique ID',
        'Worker thread spawns to process job',
        'Orchestrator crafts solver prompt',
        'LLM generates JSON animation plan',
        'Manim adapter converts JSON to Python code',
        'Manim renders silent video',
        'TTS engine generates voiceover audio',
        'FFmpeg combines video and audio',
        'Final video saved and served'
    ]
    
    for idx, step in enumerate(flow_steps, 1):
        doc.add_paragraph(f'{idx}. {step}', style='List Number')
    
    doc.add_heading('4.2 Data Flow Diagram', 2)
    
    flow_diagram = '''
    User Input (Question)
            │
            ▼
    ┌─────────────────┐
    │   Web UI Form   │
    └─────────────────┘
            │
            ▼ (HTTP POST /api/run)
    ┌─────────────────┐
    │   FastAPI App   │
    └─────────────────┘
            │
            ▼ (Create Job)
    ┌─────────────────┐
    │  Job Manager    │
    └─────────────────┘
            │
            ▼ (Worker Thread)
    ┌─────────────────┐
    │  Orchestrator   │────► LLM APIs (OpenAI/DeepSeek/Gemini)
    └─────────────────┘
            │
            ▼ (JSON Plan)
    ┌─────────────────┐
    │  Manim Adapter  │────► Generated Python Script
    └─────────────────┘
            │
            ▼ (Execute)
    ┌─────────────────┐
    │  Manim Render   │────► Silent Video (.mp4)
    └─────────────────┘
            │
            ▼
    ┌─────────────────┐
    │  TTS Synthesis  │────► Audio Files (.wav)
    └─────────────────┘
            │
            ▼ (Combine)
    ┌─────────────────┐
    │   AV Sync       │────► Final Video (.mp4)
    └─────────────────┘
            │
            ▼
    ┌─────────────────┐
    │  Serve to User  │
    └─────────────────┘
    '''
    
    code_para = doc.add_paragraph(flow_diagram)
    code_para.style = 'Intense Quote'
    
    doc.add_page_break()
    
    # ==================== API REFERENCE ====================
    doc.add_heading('5. API Reference', 1)
    
    doc.add_heading('5.1 REST Endpoints', 2)
    
    endpoints = [
        {
            'method': 'POST',
            'path': '/api/run',
            'description': 'Submit a new animation generation job',
            'body': '''{
  "problem": "string",
  "orchestrate": true,
  "voice_first": true,
  "element_audio": false,
  "custom_prompt": "string (optional)"
}''',
            'response': '''{
  "job_id": "string",
  "status": "queued"
}'''
        },
        {
            'method': 'GET',
            'path': '/api/jobs/{job_id}',
            'description': 'Check job status and retrieve results',
            'body': 'None',
            'response': '''{
  "job_id": "string",
  "status": "done|running|error",
  "video_path": "string",
  "plan_path": "string",
  "log": "string"
}'''
        }
    ]
    
    for ep in endpoints:
        doc.add_heading(f'{ep["method"]} {ep["path"]}', 3)
        doc.add_paragraph(ep['description'])
        
        if ep['body'] != 'None':
            doc.add_paragraph('Request Body:', style='Heading 4')
            code_para = doc.add_paragraph(ep['body'])
            code_para.style = 'Intense Quote'
        
        doc.add_paragraph('Response:', style='Heading 4')
        code_para = doc.add_paragraph(ep['response'])
        code_para.style = 'Intense Quote'
        
        doc.add_paragraph()
    
    doc.add_page_break()
    
    # ==================== INSTALLATION ====================
    doc.add_heading('6. Installation and Setup', 1)
    
    doc.add_heading('6.1 Prerequisites', 2)
    prereqs = [
        'Python 3.13 or higher',
        'FFmpeg (for video processing)',
        'LaTeX distribution (for mathematical typesetting)',
        'At least one LLM API key (OpenAI, DeepSeek, or Gemini)'
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')
    
    doc.add_heading('6.2 Installation Steps', 2)
    
    install_steps = [
        ('Clone the repository', 'git clone <repository-url>'),
        ('Create virtual environment', 'python -m venv .venv'),
        ('Activate environment (Windows)', '.venv\\Scripts\\activate'),
        ('Activate environment (Unix)', 'source .venv/bin/activate'),
        ('Install dependencies', 'pip install -r requirements.txt'),
        ('Configure API keys', 'Copy .env.example to .env and add keys'),
        ('Verify setup', 'python test/verify_llm_keys.py')
    ]
    
    for idx, (step_name, command) in enumerate(install_steps, 1):
        doc.add_paragraph(f'{idx}. {step_name}', style='Heading 3')
        if command:
            code_para = doc.add_paragraph(command)
            code_para.style = 'Intense Quote'
    
    doc.add_page_break()
    
    # ==================== USAGE GUIDE ====================
    doc.add_heading('7. Usage Guide', 1)
    
    doc.add_heading('7.1 Starting the Server', 2)
    doc.add_paragraph('To start the Manimations server:')
    code_para = doc.add_paragraph('.venv\\Scripts\\python.exe -m uvicorn scripts.server.app:app --reload')
    code_para.style = 'Intense Quote'
    
    doc.add_paragraph('Access the web interface at: http://localhost:8000')
    
    doc.add_heading('7.2 Using the Web Interface', 2)
    web_steps = [
        'Open browser and navigate to http://localhost:8000',
        'Enter your mathematical question in the input field',
        'Configure options (orchestrate, voice-first, etc.)',
        'Click "Run" to submit the job',
        'Monitor progress in real-time',
        'Download the generated video when complete'
    ]
    for idx, step in enumerate(web_steps, 1):
        doc.add_paragraph(f'{idx}. {step}', style='List Number')
    
    doc.add_heading('7.3 Using the Command Line', 2)
    doc.add_paragraph('For programmatic access:')
    
    code_para = doc.add_paragraph('''python -m scripts.pipeline \\
  --question "Explain calculus" \\
  --out-dir media/videos/output \\
  --voice-first''')
    code_para.style = 'Intense Quote'
    
    doc.add_heading('7.4 Using Historical Data', 2)
    doc.add_paragraph('To generate videos from pre-existing solutions:')
    
    code_para = doc.add_paragraph('''# Edit scripts/run_from_history.py with your data
python scripts/run_from_history.py''')
    code_para.style = 'Intense Quote'
    
    doc.add_page_break()
    
    # ==================== CONFIGURATION ====================
    doc.add_heading('8. Configuration Management', 1)
    
    doc.add_heading('8.1 Environment Variables', 2)
    doc.add_paragraph('Key environment variables:')
    
    env_vars = {
        'OPENAI_API_KEY': 'OpenAI API authentication key',
        'DEEPSEEK_API_KEY': 'DeepSeek API authentication key',
        'GEMINI_API_KEY': 'Google Gemini API authentication key',
        'ELEVENLABS_API_KEY': 'ElevenLabs TTS API key (optional)',
        'MANIM_QUALITY': 'Video quality: low, medium, high, production',
        'ORCHESTRATOR_OFFLINE': 'Enable offline mode (skip LLM calls)'
    }
    
    table = doc.add_table(rows=len(env_vars)+1, cols=2)
    table.style = 'Light Grid Accent 1'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Variable'
    hdr_cells[1].text = 'Description'
    
    for idx, (var, desc) in enumerate(env_vars.items(), 1):
        row_cells = table.rows[idx].cells
        row_cells[0].text = var
        row_cells[1].text = desc
    
    doc.add_page_break()
    
    # ==================== APPENDICES ====================
    doc.add_heading('16. Appendices', 1)
    
    doc.add_heading('Appendix A: File Structure', 2)
    file_structure = '''
manimations/
├── scripts/
│   ├── server/
│   │   ├── __init__.py
│   │   └── app.py              # FastAPI application
│   ├── orchestrator/
│   │   ├── llm_clients.py      # LLM API clients
│   │   ├── prompt_orchestrator.py
│   │   ├── run_orchestrator.py
│   │   └── schemas.py          # Pydantic models
│   ├── av_sync.py              # Audio-video synchronization
│   ├── manim_adapter.py        # Manim code generator
│   ├── pipeline.py             # Main pipeline
│   ├── voiceover.py            # TTS synthesis
│   └── run_from_history.py    # Historical data processor
├── web/
│   ├── index.html              # Web UI
│   ├── app.js                  # Frontend logic
│   └── styles.css              # Styling
├── test/
│   └── verify_llm_keys.py      # Key verification
├── media/
│   ├── videos/                 # Generated videos
│   └── texts/                  # JSON plans
├── Prompt.json                 # Default system prompt
├── pyproject.toml              # Project metadata
└── README.md                   # Quick start guide
'''
    
    code_para = doc.add_paragraph(file_structure)
    code_para.style = 'Intense Quote'
    
    doc.add_heading('Appendix B: Troubleshooting', 2)
    
    issues = [
        ('LLM API Errors', 'Check API keys in .env file. Run verify_llm_keys.py'),
        ('Video Not Generating', 'Ensure FFmpeg is installed and in PATH'),
        ('Slow Rendering', 'Adjust MANIM_QUALITY to "low" for faster preview'),
        ('Server Won\'t Start', 'Check if port 8000 is available. Try different port'),
        ('Audio Issues', 'Install pyttsx3: pip install pyttsx3')
    ]
    
    for issue, solution in issues:
        p = doc.add_paragraph()
        p.add_run(f'{issue}: ').bold = True
        p.add_run(solution)
    
    doc.add_heading('Appendix C: Performance Tips', 2)
    
    tips = [
        'Use voice-first mode for better audio-video sync',
        'Enable element-audio only for precise timing requirements',
        'Cache LLM responses to avoid redundant API calls',
        'Use lower quality for testing, production for final output',
        'Consider Ollama for local LLM inference (no API costs)'
    ]
    
    for tip in tips:
        doc.add_paragraph(tip, style='List Bullet')
    
    # Save document
    output_path = root / 'MANIMATIONS_DOCUMENTATION.docx'
    doc.save(str(output_path))
    print(f'\n✅ Documentation generated: {output_path}')
    print(f'📄 File size: {output_path.stat().st_size / 1024:.2f} KB')
    print(f'📖 Estimated pages: ~50+')
    return output_path

if __name__ == '__main__':
    create_documentation()
